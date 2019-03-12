'''
    功能：将字典转化为pdf
    所需的库：docx, win32com
    安装命令：pip install python-docx
             pip install pypiwin32
'''
from docx import Document
from docx.shared import Inches, Pt
from win32com.client import Dispatch, constants, gencache
import sys
import os

'''
    以下变量为与页面上选择按键的接口，
    顺序与字典中一一对应，
    值为1表示选中，值为0 表示未选中
'''
button_author = 1
button_title = 1
button_date = 1
button_text = 1
button_others = 1
show = [button_author, button_title, button_date, button_text, button_others]

'''
    需要出现在报表中的数据，字典形式储存
'''
data = [
    { 
        'Author': 'Staraptor',
        'title': 'Good Good Study',
        'published_date': '2017',
        'text': 'Too Young',
        'others': '改革春风吹满地',
    },
    {
        'Author': 'HIT',
        'title': 'Day Day Up',
        'published_date': '2018',
        'text': 'Too Simple',
        'others': '中国人民真争气',
    },
    {
        'Author': 'HITSZ',
        'title': '2333333',
        'published_date': '2019',
        'text': 'Sometimes Naive',
        'others': '好！',
    },
    ]

'''
    函数功能：创建word文件，编辑数据
'''
def creat_word():
    try:
        doc = Document('report.docx')
    except Exception:   #当文件不存在时新建
        doc = Document()
        doc.save('report.docx')
        doc = Document('report.docx')
        #插入标题
        doc.add_heading('数据报表', 0)
        p = doc.add_paragraph('')
        p.add_run('哈尔滨工业大学').font.size = Pt(12)
        p.add_run('Harbin Institude of technology').bold = True
        #插入文字

        paper_number = 0    # 论文篇数
        for dic in data:    # 遍历列表内的字典
            paper_number += 1
            doc.add_heading('论文' + str(paper_number), level = 1)
            i = 0       # 当前字典内的键值数
            for key in dic:     # 遍历字典内的键值
                if (show[i] == 1):    # 逐条打印已选择的内容
                    doc.add_paragraph(key + ' : ' + dic[key])
                i = i + 1
        #插入图片
        images = 'sakura.jpg'
        doc.add_picture(images, width = Inches(4))
        #保存word文件
        doc.save('report.docx')

    else:   #文件存在时直接向里面插入编辑
        doc.add_heading('数据报表', 0)
        p = doc.add_paragraph('')
        p.add_run('哈尔滨工业大学').font.size = Pt(12)
        p.add_run('Harbin Institude of technology').bold = True

        paper_number = 0
        for dic in data:
            paper_number += 1 
            doc.add_heading('论文' + str(paper_number), level = 1)
            i = 0
            for key in dic:
                if (show[i] == 1):
                    doc.add_paragraph(key + ' : ' + dic[key])
                i = i+1
        images = 'sakura.jpg'
        doc.add_picture(images, width = Inches(4))
        doc.save('report.docx')

'''
    函数功能：将word转化为pdf，并调用系统路径保存在用户桌面
'''
def creat_report():
    #print(len(sys.argv))
    if (len(sys.argv) == 2):
        input = sys.argv[1]
        output = os.path.splitext(input)[0]+'.pdf'
    elif (len(sys.argv) == 3):
        input = sys.argv[1]
        output = sys.argv[2]
    else:
        #调用os库使报表生成在用户桌面
        desktop = os.path.join(os.path.expanduser("~"), 'Desktop')
        desktop_path = desktop + './report.pdf'
        input = u'report.docx'#word文档的名称
        output = desktop_path #pdf文档的名称
    if (not os.path.isabs(input)):
        input = os.path.abspath(input)
    if (not os.path.isabs(output)):
        output = os.path.abspath(output)
    try:
        GenerateSupport()
        doc2pdf(input, output)
        print('Creat report Successfully!') #生成报表成功
        os.remove("report.docx") #生成pdf完毕后，删除原本的word文件
    except:
        print('Creat report Fail!') #生成报表失败

'''
    函数功能：作为文件路径的输入输出函数
'''
def doc2pdf(input, output):
    w = Dispatch("Word.Application")
    try:
        doc = w.Documents.Open(input, ReadOnly = 1)
        doc.ExportAsFixedFormat(output, constants.wdExportFormatPDF,\
            Item = constants.wdExportDocumentWithMarkup, CreateBookmarks = constants.wdExportCreateHeadingBookmarks)
        return 0
    except:
        return 1
    finally:
        w.Quit(constants.wdDoNotSaveChanges)

'''
    函数功能：解决老版本word转化问题
'''
def GenerateSupport():
    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)

'''
    函数功能：将创建word和转为pdf综合为一个函数
'''
def Run_Creat_Report():
    creat_word()
    creat_report() 

if __name__ == "__main__":
    Run_Creat_Report()


        
